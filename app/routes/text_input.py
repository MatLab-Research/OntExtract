from flask import Blueprint, render_template, request, flash, redirect, url_for, jsonify, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
import os
import uuid
from datetime import datetime
from docx import Document as DocxDocument
from langdetect import detect

from app import db
from app.models.document import Document
from app.models.processing_job import ProcessingJob
from sqlalchemy import text
from app.services.text_processing import TextProcessingService
from app.utils.file_handler import FileHandler
from app.utils.auth_decorators import write_login_required, public_with_auth_context, api_require_login_for_write

text_input_bp = Blueprint('text_input', __name__)

@text_input_bp.route('/')
@text_input_bp.route('/upload')
@write_login_required  # Only require login for upload
def upload_form():
    """Main upload form page with enhanced metadata extraction"""
    return render_template('text_input/upload_enhanced.html')

@text_input_bp.route('/paste')
@write_login_required  # Only require login for paste
def paste_form():
    """Text paste form page"""
    return render_template('text_input/paste.html')

@text_input_bp.route('/submit_text', methods=['POST'])
@write_login_required  # Require login for submission
def submit_text():
    """Handle pasted text submission"""
    try:
        data = request.get_json() if request.is_json else request.form
        
        # Get form data
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        
        if not content:
            if request.is_json:
                return jsonify({'error': 'Content is required'}), 400
            flash('Content is required', 'error')
            return redirect(url_for('text_input.paste_form'))
        
        # Auto-generate title if not provided
        if not title:
            first_line = content.split('\n')[0].strip()
            title = first_line[:50] + ('...' if len(first_line) > 50 else '') or 'Untitled Text'
        
        # Detect language
        detected_language = None
        language_confidence = 0.0
        try:
            detected_language = detect(content)
            language_confidence = 0.9  # Langdetect doesn't provide confidence, using default
        except:
            detected_language = 'en'  # Default to English
            language_confidence = 0.5
        
        # Create document record
        document = Document(
            title=title,
            content_type='text',
            content=content,
            detected_language=detected_language,
            language_confidence=language_confidence,
            status='uploaded',
            user_id=current_user.id
        )

        db.session.add(document)
        db.session.commit()

        # Track document upload with PROV-O
        try:
            from app.services.provenance_service import provenance_service
            provenance_service.track_document_upload(document, current_user)
        except Exception as e:
            current_app.logger.warning(f"Failed to track document upload provenance: {str(e)}")

        # Note: Segmentation is now manual from document processing page
        # Removed automatic segmentation to allow user control

        if request.is_json:
            return jsonify({
                'success': True,
                'document_id': document.id,
                'message': 'Text submitted successfully',
                'redirect_url': url_for('text_input.document_detail', document_uuid=document.uuid)
            })

        flash('Text submitted successfully!', 'success')
        return redirect(url_for('text_input.document_detail', document_uuid=document.uuid))
        
    except Exception as e:
        current_app.logger.error(f"Error submitting text: {str(e)}")
        if request.is_json:
            return jsonify({'error': 'An error occurred while processing your text'}), 500
        flash('An error occurred while processing your text', 'error')
        return redirect(url_for('text_input.paste_form'))

@text_input_bp.route('/upload_file', methods=['POST'])
@write_login_required  # Require login for upload
def upload_file():
    """Handle file upload submission"""
    try:
        if 'file' not in request.files:
            if request.is_json:
                return jsonify({'error': 'No file selected'}), 400
            flash('No file selected', 'error')
            return redirect(url_for('text_input.upload_form'))
        
        file = request.files['file']
        title = request.form.get('title', '').strip()
        
        if file.filename == '':
            if request.is_json:
                return jsonify({'error': 'No file selected'}), 400
            flash('No file selected', 'error')
            return redirect(url_for('text_input.upload_form'))
        
        # Validate file type
        file_handler = FileHandler()
        fname = file.filename or ''
        if not file_handler.allowed_file(fname):
            allowed = ', '.join(current_app.config['ALLOWED_EXTENSIONS'])
            error_msg = f'File type not allowed. Allowed types: {allowed}'
            if request.is_json:
                return jsonify({'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('text_input.upload_form'))
        
        # Generate secure filename
        original_filename = fname
        filename = secure_filename(original_filename or '')
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        
        # Save file
        upload_folder = current_app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, unique_filename)
        file.save(file_path)
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Extract text content
        content = file_handler.extract_text_from_file(file_path, original_filename or '')
        
        if not content:
            os.remove(file_path)  # Clean up file
            error_msg = 'Could not extract text from file'
            if request.is_json:
                return jsonify({'error': error_msg}), 400
            flash(error_msg, 'error')
            return redirect(url_for('text_input.upload_form'))
        
        # Auto-generate title if not provided
        if not title:
            title = os.path.splitext(original_filename or '')[0]
        
        # Detect language
        detected_language = None
        language_confidence = 0.0
        try:
            detected_language = detect(content)
            language_confidence = 0.9
        except:
            detected_language = 'en'
            language_confidence = 0.5
        
        # Get file extension
        file_extension = file_handler.get_file_extension(original_filename or '')
        
        # Create document record
        document = Document(
            title=title,
            content_type='file',
            file_type=file_extension,
            original_filename=original_filename,
            file_path=file_path,
            file_size=file_size,
            content=content,
            detected_language=detected_language,
            language_confidence=language_confidence,
            status='uploaded',
            user_id=current_user.id
        )

        db.session.add(document)
        db.session.commit()

        # Track document upload with PROV-O
        try:
            from app.services.provenance_service import provenance_service
            provenance_service.track_document_upload(document, current_user)
        except Exception as e:
            current_app.logger.warning(f"Failed to track document upload provenance: {str(e)}")

        # Note: Segmentation is now manual from document processing page
        # Removed automatic segmentation to allow user control

        if request.is_json:
            return jsonify({
                'success': True,
                'document_id': document.id,
                'message': 'File uploaded successfully',
                'redirect_url': url_for('text_input.document_detail', document_uuid=document.uuid)
            })

        flash('File uploaded successfully!', 'success')
        return redirect(url_for('text_input.document_detail', document_uuid=document.uuid))
        
    except Exception as e:
        current_app.logger.error(f"Error uploading file: {str(e)}")
        # Clean up file if it was saved
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        
        if request.is_json:
            return jsonify({'error': 'An error occurred while uploading your file'}), 500
        flash('An error occurred while uploading your file', 'error')
        return redirect(url_for('text_input.upload_form'))

@text_input_bp.route('/documents')
@public_with_auth_context  # Allow anonymous viewing of document list
def document_list():
    """List all documents grouped by base document with version stacking - public access"""
    # Show all documents for everyone (public access)
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    # Get all documents ordered by creation date (newest first)
    all_documents = Document.query.order_by(Document.created_at.desc()).all()
    
    # Group documents by base document
    document_groups = {}
    for doc in all_documents:
        # Determine the base document ID
        base_id = doc.source_document_id or doc.id
        
        if base_id not in document_groups:
            document_groups[base_id] = {
                'base_document': None,
                'versions': [],
                'latest_created': None
            }
        
        # Set the base document (original version)
        if doc.version_type == 'original':
            document_groups[base_id]['base_document'] = doc
        
        # Add to versions list
        document_groups[base_id]['versions'].append(doc)
        
        # Track latest creation date for sorting groups
        if document_groups[base_id]['latest_created'] is None or doc.created_at > document_groups[base_id]['latest_created']:
            document_groups[base_id]['latest_created'] = doc.created_at
    
    # Sort versions within each group (latest first)
    for group in document_groups.values():
        group['versions'].sort(key=lambda x: x.version_number, reverse=True)
        # If no base document was found, use the original version
        if group['base_document'] is None and group['versions']:
            group['base_document'] = min(group['versions'], key=lambda x: x.version_number)
    
    # Convert to list and sort by latest creation date
    grouped_documents = list(document_groups.values())
    grouped_documents.sort(key=lambda x: x['latest_created'], reverse=True)
    
    # Implement pagination on groups
    start = (page - 1) * per_page
    end = start + per_page
    paginated_groups = grouped_documents[start:end]
    
    # Create pagination object
    total_groups = len(grouped_documents)
    has_prev = page > 1
    has_next = end < total_groups
    
    pagination = type('Pagination', (), {
        'items': paginated_groups,
        'page': page,
        'pages': (total_groups + per_page - 1) // per_page,
        'has_prev': has_prev,
        'has_next': has_next,
        'prev_num': page - 1 if has_prev else None,
        'next_num': page + 1 if has_next else None,
        'iter_pages': lambda: range(1, (total_groups + per_page - 1) // per_page + 1)
    })()
    
    return render_template('text_input/document_list.html', documents=pagination)

@text_input_bp.route('/document/<uuid:document_uuid>')
@public_with_auth_context  # Allow anonymous viewing of documents
def document_detail(document_uuid):
    """Show document details - simplified experiment-centric view"""
    # Get document - public access for viewing
    document = Document.query.filter_by(uuid=document_uuid).first_or_404()

    # Get experiments that include this document with their processing results
    from app.models.experiment_document import ExperimentDocument
    from app.models.experiment_processing import DocumentProcessingIndex

    # Get all experiment-document relationships for this document
    experiment_documents = ExperimentDocument.query.filter_by(document_id=document.id).all()

    # Enrich with processing information
    document_experiments = []
    total_processing_count = 0

    for exp_doc in experiment_documents:
        # Get processing operations for this experiment-document pair
        processing_results = DocumentProcessingIndex.query.filter_by(
            document_id=document.id,
            experiment_id=exp_doc.experiment_id
        ).all()

        # Create a data structure for the template
        exp_data = {
            'experiment': exp_doc.experiment,
            'processing_results': [
                {
                    'processing_type': proc.processing_type,
                    'processing_method': proc.processing_method,
                    'status': proc.status
                }
                for proc in processing_results
            ]
        }
        document_experiments.append(exp_data)
        total_processing_count += len(processing_results)

    return render_template('text_input/document_detail_simplified.html',
                         document=document,
                         document_experiments=document_experiments,
                         total_processing_count=total_processing_count)

@text_input_bp.route('/document/<uuid:document_uuid>/delete', methods=['POST'])
@write_login_required  # Require login for deletion
def delete_document(document_uuid):
    """Delete a document"""
    document = Document.query.filter_by(uuid=document_uuid).first_or_404()
    
    try:
        # Delete associated file
        document.delete_file()
        
        # Delete database record (cascades to related records)
        db.session.delete(document)
        db.session.commit()
        
        if request.is_json:
            return jsonify({'success': True, 'message': 'Document deleted successfully'})
        
        flash('Document deleted successfully', 'success')
        return redirect(url_for('text_input.document_list'))
        
    except Exception as e:
        current_app.logger.error(f"Error deleting document: {str(e)}")
        
        if request.is_json:
            return jsonify({'error': 'An error occurred while deleting the document'}), 500
        
        flash('An error occurred while deleting the document', 'error')
        return redirect(url_for('text_input.document_detail', document_uuid=document_uuid))


@text_input_bp.route('/document/<uuid:document_uuid>/metadata', methods=['GET'])
@public_with_auth_context
def get_document_metadata(document_uuid):
    """Get document metadata for editing"""
    document = Document.query.filter_by(uuid=document_uuid).first_or_404()

    try:
        metadata = document.source_metadata or {}

        # Include title from document.title field (may differ from source_metadata)
        if 'title' not in metadata or not metadata.get('title'):
            metadata['title'] = document.title

        # Handle authors formatting
        if 'authors' in metadata:
            if isinstance(metadata['authors'], list):
                metadata['authors'] = ', '.join(metadata['authors'])

        return jsonify({
            'success': True,
            'metadata': metadata
        })
    except Exception as e:
        current_app.logger.error(f"Error fetching metadata: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@text_input_bp.route('/document/<uuid:document_uuid>/metadata', methods=['PUT'])
@write_login_required
def update_document_metadata(document_uuid):
    """Update document metadata with provenance tracking"""
    document = Document.query.filter_by(uuid=document_uuid).first_or_404()

    try:
        new_data = request.get_json()

        # Get old metadata
        old_metadata = document.source_metadata or {}

        # Track changes for provenance
        changes = {}

        # Build new metadata and track changes
        new_metadata = {}
        fields = ['title', 'authors', 'publication_date', 'journal', 'publisher',
                  'type', 'doi', 'isbn', 'url', 'abstract', 'citation']

        for field in fields:
            new_value = new_data.get(field, '').strip()
            old_value = old_metadata.get(field, '')

            # Convert authors string to list if provided
            if field == 'authors' and new_value:
                new_value = [a.strip() for a in new_value.split(',') if a.strip()]
                if isinstance(old_value, list):
                    old_value_str = ', '.join(old_value)
                else:
                    old_value_str = old_value
                new_value_str = ', '.join(new_value)

                if old_value_str != new_value_str:
                    changes[field] = {'old': old_value, 'new': new_value}
            elif field == 'title':
                # Update document title field as well
                if new_value and new_value != document.title:
                    document.title = new_value
                if str(old_value) != str(new_value):
                    changes[field] = {'old': old_value, 'new': new_value}
            else:
                # Track changes for other fields
                if str(old_value) != str(new_value):
                    changes[field] = {'old': old_value, 'new': new_value}

            # Only add non-empty values to new metadata
            if new_value:
                new_metadata[field] = new_value

        # Update document metadata
        document.source_metadata = new_metadata
        document.updated_at = datetime.utcnow()

        db.session.commit()

        # Track metadata update with PROV-O if there were changes
        if changes:
            try:
                from app.services.provenance_service import provenance_service
                provenance_service.track_metadata_update(document, current_user, changes)
            except Exception as e:
                current_app.logger.warning(f"Failed to track metadata update provenance: {str(e)}")

        return jsonify({
            'success': True,
            'message': 'Metadata updated successfully',
            'changes_count': len(changes)
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating metadata: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@text_input_bp.route('/document/<int:base_document_id>/delete-all-versions', methods=['POST'])
@api_require_login_for_write
def delete_all_versions(base_document_id):
    """Delete all versions of a document family"""
    try:
        # Find the base document
        base_document = Document.query.get_or_404(base_document_id)
        
        # Get all documents in this family (base + all versions that reference it)
        all_versions = []
        
        # If this is already a processed version, find its original
        if base_document.source_document_id:
            actual_base = Document.query.get(base_document.source_document_id)
            if actual_base:
                base_document = actual_base
        
        # Add the base document
        all_versions.append(base_document)
        
        # Find all versions that reference this base document
        derived_versions = Document.query.filter_by(source_document_id=base_document.id).all()
        all_versions.extend(derived_versions)
        
        # Also find any versions that might reference the same original as this one
        # (in case base_document_id wasn't the true original)
        if base_document.source_document_id:
            sibling_versions = Document.query.filter_by(source_document_id=base_document.source_document_id).all()
            for sibling in sibling_versions:
                if sibling not in all_versions:
                    all_versions.append(sibling)
        
        deleted_count = 0
        document_title = base_document.title
        
        # Delete all documents in the family
        for document in all_versions:
            try:
                # Delete associated data first
                # Delete processing jobs
                ProcessingJob.query.filter_by(document_id=document.id).delete()
                
                # Delete text segments  
                db.session.execute(text("DELETE FROM text_segments WHERE document_id = :doc_id"), 
                                 {'doc_id': document.id})
                
                # Delete document embeddings
                db.session.execute(text("DELETE FROM document_embeddings WHERE document_id = :doc_id"), 
                                 {'doc_id': document.id})
                
                # Delete the document itself
                db.session.delete(document)
                deleted_count += 1
                
            except Exception as e:
                current_app.logger.error(f"Error deleting document {document.id}: {str(e)}")
                continue
        
        db.session.commit()
        
        current_app.logger.info(f"Successfully deleted {deleted_count} versions of document family '{document_title}'")
        
        if request.is_json:
            return jsonify({
                'success': True, 
                'deleted_count': deleted_count,
                'document_title': document_title
            })
        
        flash(f'Successfully deleted {deleted_count} document versions', 'success')
        return redirect(url_for('text_input.document_list'))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting document family {base_document_id}: {str(e)}")
        if request.is_json:
            return jsonify({'success': False, 'error': str(e)}), 500
        flash('Error deleting document versions', 'error')
        return redirect(url_for('text_input.document_list'))

@text_input_bp.route('/api/document/<int:document_id>/content')
@api_require_login_for_write
def api_document_content(document_id):
    """API endpoint to get document content"""
    document = Document.query.filter_by(id=document_id).first_or_404()
    return jsonify(document.to_dict(include_content=True))

@text_input_bp.route('/api/documents')
@api_require_login_for_write
def api_document_list():
    """API endpoint to list user's documents"""
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    
    documents = Document.query\
        .order_by(Document.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'documents': [doc.to_dict() for doc in documents.items],
        'total': documents.total,
        'pages': documents.pages,
        'current_page': documents.page,
        'has_next': documents.has_next,
        'has_prev': documents.has_prev
    })

@text_input_bp.route('/documents/<int:document_id>/apply_embeddings', methods=['POST'])
@write_login_required  # Require login for processing
def apply_embeddings(document_id):
    """Apply embeddings to a document"""
    try:
        document = Document.query.filter_by(id=document_id).first_or_404()
        
        if not document.content:
            return jsonify({'error': 'Document has no content to process'}), 400
        
        # Initialize embedding service
        try:
            from shared_services.embedding.embedding_service import EmbeddingService
            embedding_service = EmbeddingService()
        except ImportError:
            # Fallback to basic implementation if shared services not available
            return jsonify({'error': 'Embedding service not available'}), 500
        
        # Generate embeddings
        try:
            # Process document content in chunks if too long
            content = document.content
            max_length = 8000  # Conservative limit for most embedding models
            
            if len(content) > max_length:
                # Split into chunks and embed each
                chunks = [content[i:i+max_length] for i in range(0, len(content), max_length)]
                embeddings = []
                for chunk in chunks:
                    chunk_embedding = embedding_service.get_embedding(chunk)
                    embeddings.append(chunk_embedding)
                
                # Store metadata about chunked processing
                embedding_info = {
                    'type': 'chunked',
                    'chunks': len(chunks),
                    'chunk_size': max_length,
                    'model': embedding_service.get_model_name(),
                    'dimension': embedding_service.get_dimension()
                }
            else:
                # Single embedding for short documents
                embeddings = [embedding_service.get_embedding(content)]
                embedding_info = {
                    'type': 'single',
                    'model': embedding_service.get_model_name(),
                    'dimension': embedding_service.get_dimension()
                }
            
            # Update document metadata
            if not document.processing_metadata:
                document.processing_metadata = {}
            
            # Mark embeddings as applied
            document.processing_metadata['processing_info'] = document.processing_metadata.get('processing_info', {})
            document.processing_metadata['processing_info']['embeddings_applied'] = True
            document.processing_metadata['processing_info']['embeddings_info'] = embedding_info
            document.processing_metadata['processing_info']['applied_at'] = datetime.utcnow().isoformat()
            
            # Update word count if not set
            if not document.word_count:
                document.word_count = len(content.split())
            
            document.updated_at = datetime.utcnow()
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Embeddings applied successfully',
                'embedding_info': embedding_info
            })
            
        except Exception as e:
            current_app.logger.error(f"Error generating embeddings: {str(e)}")
            return jsonify({'error': f'Failed to generate embeddings: {str(e)}'}), 500
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error applying embeddings: {str(e)}")
        return jsonify({'error': 'An error occurred while applying embeddings'}), 500


# Composite Document Routes

@text_input_bp.route('/composite/create/<int:document_id>', methods=['POST'])
@write_login_required
def create_composite(document_id):
    """Create a composite document from all versions of the given document"""
    try:
        # Get the document and all its versions
        document = Document.query.get_or_404(document_id)
        all_versions = document.get_all_versions()

        # Filter to get only processed versions (exclude original and composites)
        processed_versions = [v for v in all_versions if v.version_type == 'processed']

        if len(processed_versions) < 2:
            return jsonify({
                'success': False,
                'error': 'Need at least 2 processed versions to create a composite'
            }), 400

        # Import the composite service
        from app.services.composite_versioning_service import CompositeVersioningService

        # Create the composite
        composite_doc = CompositeVersioningService.create_composite_from_versions(
            original_document=document.get_root_document(),
            source_versions=processed_versions
        )

        return jsonify({
            'success': True,
            'composite_id': composite_doc.id,
            'message': f'Composite created with {len(processed_versions)} source versions'
        })

    except Exception as e:
        current_app.logger.error(f"Error creating composite: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@text_input_bp.route('/composite/sources/<int:composite_id>')
@public_with_auth_context
def get_composite_sources(composite_id):
    """Get information about the source documents for a composite"""
    try:
        composite = Document.query.get_or_404(composite_id)

        if not composite.is_composite():
            return jsonify({
                'success': False,
                'error': 'Document is not a composite'
            }), 400

        sources = composite.get_composite_sources()

        source_data = []
        for source in sources:
            source_data.append({
                'id': source.id,
                'title': source.title,
                'version_type': source.version_type,
                'version_number': source.version_number,
                'created_at': source.created_at.isoformat() if source.created_at else None,
                'processing_metadata': source.processing_metadata
            })

        return jsonify({
            'success': True,
            'composite_id': composite_id,
            'source_documents': source_data
        })

    except Exception as e:
        current_app.logger.error(f"Error getting composite sources: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@text_input_bp.route('/composite/update/<int:composite_id>', methods=['POST'])
@write_login_required
def update_composite(composite_id):
    """Update a composite document with new versions"""
    try:
        composite = Document.query.get_or_404(composite_id)

        if not composite.is_composite():
            return jsonify({
                'success': False,
                'error': 'Document is not a composite'
            }), 400

        # Get the root document and find any new versions
        root_doc = composite.get_root_document()
        all_versions = root_doc.get_all_versions()
        existing_sources = set(composite.composite_sources or [])

        # Find new processed versions not in the composite
        new_versions = [
            v for v in all_versions
            if v.version_type == 'processed' and v.id not in existing_sources
        ]

        if not new_versions:
            return jsonify({
                'success': False,
                'error': 'No new versions to add to composite'
            }), 400

        # Import the composite service
        from app.services.composite_versioning_service import CompositeVersioningService

        # Update the composite with new versions
        for new_version in new_versions:
            CompositeVersioningService.update_composite_from_new_version(composite, new_version)

        return jsonify({
            'success': True,
            'message': f'Added {len(new_versions)} new versions to composite'
        })

    except Exception as e:
        current_app.logger.error(f"Error updating composite: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# Enhanced Upload Routes with Metadata Extraction and Provenance

print("=== REGISTERING /upload/extract_metadata ROUTE ===", flush=True)

@text_input_bp.route('/upload/extract_metadata', methods=['POST'])
@write_login_required
def extract_upload_metadata():
    """Extract metadata from uploaded document or DOI"""
    from datetime import datetime
    # Write to debug file
    with open('/tmp/ontextract_debug.log', 'a') as f:
        f.write(f"\n===== FUNCTION CALLED at {datetime.now()} =====\n")
        f.flush()

    print("=== INSIDE FUNCTION - START ===", flush=True)
    from app.services.upload_service import upload_service
    import json

    print("=== EXTRACT_METADATA CALLED ===", flush=True)
    current_app.logger.error("=== EXTRACT_METADATA CALLED (ERROR LEVEL) ===")

    try:
        extraction_source = request.form.get('source_type')
        print(f"Extraction source: {extraction_source}", flush=True)

        if extraction_source == 'doi':
            # Extract from DOI using shared service
            doi = request.form.get('doi')
            if not doi:
                return jsonify({'error': 'DOI is required'}), 400

            # Get bibliographic metadata from CrossRef
            metadata_result = upload_service.extract_metadata_from_doi(doi)

            if not metadata_result.success:
                return jsonify({'error': metadata_result.error}), 404

            # Build provenance tracking
            provenance = {}
            for key, value in metadata_result.metadata.items():
                if value is not None:
                    provenance[key] = {
                        'source': 'crossref',
                        'confidence': 0.95,
                        'timestamp': datetime.utcnow().isoformat(),
                        'raw_value': value
                    }

            # Return metadata for review (no file yet - user must upload)
            return jsonify({
                'success': True,
                'metadata': metadata_result.metadata,
                'provenance': provenance,
                'needs_file': True,
                'message': 'Bibliographic metadata retrieved. Please upload the document file.'
            })

        elif extraction_source == 'file':
            # Process uploaded file
            if 'document_file' not in request.files:
                return jsonify({'error': 'No file uploaded'}), 400

            file = request.files['document_file']

            # Save to temporary location
            upload_result = upload_service.save_to_temp(file)
            if not upload_result.success:
                return jsonify({'error': upload_result.error}), 400

            try:
                crossref_metadata = {}
                crossref_provenance = {}

                # Get user-provided title (if any)
                user_title = request.form.get('title', '').strip()
                current_app.logger.info(f"Upload metadata extraction starting. User title: '{user_title}', File: {file.filename}")

                # Try automatic PDF analysis first if no user title provided
                if not user_title and file.filename.lower().endswith('.pdf'):
                    current_app.logger.info("Analyzing PDF for DOI/title...")
                    pdf_result = upload_service.extract_metadata_from_pdf(upload_result.temp_path)

                    if pdf_result.success:
                        crossref_metadata = pdf_result.metadata
                        # Track CrossRef provenance with note about automatic extraction
                        for key, value in crossref_metadata.items():
                            if value is not None and key not in ['extracted_doi', 'extracted_title', 'extraction_method']:
                                crossref_provenance[key] = {
                                    'source': 'crossref_auto',
                                    'confidence': 0.90,
                                    'timestamp': datetime.utcnow().isoformat(),
                                    'raw_value': value,
                                    'extraction_method': pdf_result.metadata.get('extraction_method', 'auto')
                                }
                        current_app.logger.info(f"Automatic extraction successful: {pdf_result.metadata.get('extraction_method')}")
                    else:
                        current_app.logger.warning(f"Automatic extraction failed: {pdf_result.error if hasattr(pdf_result, 'error') else 'Unknown error'}")

                # If user provided title, use that (takes precedence)
                if user_title:
                    current_app.logger.info(f"Using user-provided title: {user_title}")
                    metadata_result = upload_service.extract_metadata_from_title(user_title)
                    if metadata_result.success:
                        crossref_metadata = metadata_result.metadata
                        # Track CrossRef provenance
                        for key, value in crossref_metadata.items():
                            if value is not None:
                                crossref_provenance[key] = {
                                    'source': 'crossref',
                                    'confidence': metadata_result.metadata.get('match_score', 0.85),
                                    'timestamp': datetime.utcnow().isoformat(),
                                    'raw_value': value
                                }

                # Parse user-provided metadata
                user_metadata = {}
                user_provenance = {}

                if user_title:
                    user_metadata['title'] = user_title
                    user_provenance['title'] = {
                        'source': 'user',
                        'confidence': 1.0,
                        'timestamp': datetime.utcnow().isoformat(),
                        'raw_value': user_title
                    }

                pub_year = request.form.get('publication_year', '').strip()
                if pub_year:
                    user_metadata['publication_year'] = int(pub_year)
                    user_provenance['publication_year'] = {
                        'source': 'user',
                        'confidence': 1.0,
                        'timestamp': datetime.utcnow().isoformat(),
                        'raw_value': int(pub_year)
                    }

                authors_str = request.form.get('authors', '').strip()
                if authors_str:
                    authors = [a.strip() for a in authors_str.split(',')]
                    user_metadata['authors'] = authors
                    user_provenance['authors'] = {
                        'source': 'user',
                        'confidence': 1.0,
                        'timestamp': datetime.utcnow().isoformat(),
                        'raw_value': authors
                    }

                # Log metadata before merge
                current_app.logger.info(f"CrossRef metadata: {list(crossref_metadata.keys())}")
                current_app.logger.info(f"User metadata: {list(user_metadata.keys())}")

                # Merge metadata (user takes precedence over CrossRef)
                merged_metadata = upload_service.merge_metadata(
                    crossref_metadata,
                    user_metadata,
                    {'filename': file.filename}
                )

                current_app.logger.info(f"Merged metadata: {list(merged_metadata.keys())}")

                # Merge provenance (user takes precedence)
                merged_provenance = {**crossref_provenance, **user_provenance}

                # Add filename provenance
                merged_provenance['filename'] = {
                    'source': 'file',
                    'confidence': 1.0,
                    'timestamp': datetime.utcnow().isoformat(),
                    'raw_value': file.filename
                }

                return jsonify({
                    'success': True,
                    'metadata': merged_metadata,
                    'provenance': merged_provenance,
                    'temp_path': upload_result.temp_path,
                    'needs_file': False,
                    'message': 'Document analyzed. Please review metadata before saving.'
                })

            except Exception as e:
                # Clean up temp file
                upload_service.cleanup_temp(upload_result.temp_path)
                raise e

        else:
            return jsonify({'error': 'Invalid source type'}), 400

    except Exception as e:
        print(f"=== EXCEPTION IN EXTRACT_METADATA: {str(e)} ===", flush=True)
        import traceback
        traceback.print_exc()
        current_app.logger.error(f"Error extracting metadata: {str(e)}")
        return jsonify({'error': str(e)}), 500


@text_input_bp.route('/upload/save_document', methods=['POST'])
@write_login_required
def save_upload_document():
    """Save document after metadata review"""
    from app.services.upload_service import upload_service
    import json

    try:
        data = request.get_json()

        # Get metadata and provenance from request
        metadata = data.get('metadata', {})
        provenance = data.get('provenance', {})
        temp_path = data.get('temp_path')
        filename = data.get('filename')

        if not temp_path:
            return jsonify({'error': 'No document file to save'}), 400

        # Create upload directory
        upload_dir = current_app.config.get('UPLOAD_FOLDER', 'uploads')

        # Save permanently
        save_result = upload_service.save_permanent(temp_path, upload_dir, filename)
        if not save_result.success:
            return jsonify({'error': save_result.error}), 400

        final_path = save_result.file_path

        # Extract text content for the document
        content, error = upload_service.extract_text_content(final_path, filename)
        if error:
            return jsonify({'error': error}), 400

        # Prepare source_metadata (bibliographic info)
        source_metadata = {
            'authors': metadata.get('authors', []),
            'publication_year': metadata.get('publication_year'),
            'journal': metadata.get('journal'),
            'publisher': metadata.get('publisher'),
            'doi': metadata.get('doi'),
            'url': metadata.get('url'),
            'abstract': metadata.get('abstract'),
            'type': metadata.get('type'),
            'extraction_source': 'enhanced_upload'
        }

        # Create document record
        document = Document(
            title=metadata.get('title', filename),
            content_type='file',
            file_type='pdf',
            original_filename=filename,
            file_path=final_path,
            file_size=os.path.getsize(final_path),
            content=content,
            source_metadata=source_metadata,
            metadata_provenance=provenance,
            status='uploaded',
            user_id=current_user.id
        )

        db.session.add(document)
        db.session.commit()

        # Track document upload with PROV-O
        try:
            from app.services.provenance_service import provenance_service
            provenance_service.track_document_upload(document, current_user)

            # Track metadata extraction separately if CrossRef/Zotero was used
            if provenance:
                # Collect fields that came from automated extraction
                crossref_fields = {}
                zotero_fields = {}

                for field_name, prov_data in provenance.items():
                    if isinstance(prov_data, dict):
                        source = prov_data.get('source', '')
                        if source in ['crossref', 'crossref_auto']:
                            crossref_fields[field_name] = prov_data.get('raw_value')
                        elif source == 'zotero':
                            zotero_fields[field_name] = prov_data.get('raw_value')

                # Track CrossRef extraction if any fields were extracted
                if crossref_fields:
                    confidence = provenance.get('match_score', {}).get('raw_value', 0.9) if 'match_score' in provenance else 0.9
                    provenance_service.track_metadata_extraction(
                        document, current_user, 'crossref', crossref_fields, confidence
                    )

                # Track Zotero extraction if any fields were extracted
                if zotero_fields:
                    provenance_service.track_metadata_extraction(
                        document, current_user, 'zotero', zotero_fields, 0.95
                    )
        except Exception as e:
            current_app.logger.warning(f"Failed to track document upload provenance: {str(e)}")

        return jsonify({
            'success': True,
            'message': 'Document saved successfully',
            'document_id': document.id,
            'document_uuid': str(document.uuid)
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving document: {str(e)}")
        return jsonify({'error': str(e)}), 500
