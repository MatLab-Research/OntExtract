"""
File processing service for extracting text from various file types.

Supports:
- PDF files
- DOCX documents  
- HTML files
- Plain text files
- URLs/web pages
"""

import os
import requests
import logging
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)

class BaseFileProcessor(ABC):
    """Abstract base class for file processors."""
    
    @abstractmethod
    def process(self, file_path: str) -> str:
        """Extract text from file."""
        pass
    
    @abstractmethod
    def can_process(self, file_type: str) -> bool:
        """Check if processor can handle file type."""
        pass

class PDFProcessor(BaseFileProcessor):
    """PDF file processor using pypdf (successor of PyPDF2)."""
    
    def process(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
            
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                text = ""
                
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text += page_text + "\n\n"
                
                return text.strip()
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Install with 'pip install pypdf'")
        except Exception as e:
            raise ValueError(f"Error processing PDF {file_path}: {str(e)}")
    
    def can_process(self, file_type: str) -> bool:
        """Check if can process PDF files."""
        return file_type.lower() == 'pdf'

class DOCXProcessor(BaseFileProcessor):
    """DOCX file processor using python-docx."""
    
    def process(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with 'pip install python-docx'")
        except Exception as e:
            raise ValueError(f"Error processing DOCX {file_path}: {str(e)}")
    
    def can_process(self, file_type: str) -> bool:
        """Check if can process DOCX files."""
        return file_type.lower() == 'docx'

class HTMLProcessor(BaseFileProcessor):
    """HTML file processor using BeautifulSoup."""
    
    def process(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Get text
                text = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
        except ImportError:
            raise ImportError("BeautifulSoup4 is required for HTML processing. Install with 'pip install beautifulsoup4'")
        except Exception as e:
            raise ValueError(f"Error processing HTML {file_path}: {str(e)}")
    
    def can_process(self, file_type: str) -> bool:
        """Check if can process HTML files."""
        return file_type.lower() in ['html', 'htm']

class TextProcessor(BaseFileProcessor):
    """Plain text file processor."""
    
    def process(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error processing text file {file_path}: {str(e)}")
    
    def can_process(self, file_type: str) -> bool:
        """Check if can process text files."""
        return file_type.lower() in ['txt', 'text', 'md', 'markdown']

class URLProcessor(BaseFileProcessor):
    """URL/web page processor using BeautifulSoup."""
    
    def process(self, url: str) -> str:
        """Extract text from URL."""
        try:
            from bs4 import BeautifulSoup
            import re
            
            # Fetch URL content
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(["script", "style", "head", "meta", "noscript"]):
                element.extract()
            
            # Process and extract text while preserving structure
            result_text = []
            
            # Handle headings
            for heading in soup.find_all(re.compile('^h[1-6]$')):
                heading_text = heading.get_text().strip()
                if heading_text:
                    result_text.append(f"\n{'#' * int(heading.name[1])} {heading_text}\n")
            
            # Handle lists
            for list_element in soup.find_all(['ul', 'ol']):
                for i, item in enumerate(list_element.find_all('li'), 1):
                    item_text = item.get_text().strip()
                    if list_element.name == 'ul':
                        result_text.append(f"• {item_text}")
                    else:
                        result_text.append(f"{i}. {item_text}")
            
            # Handle paragraphs
            for paragraph in soup.find_all(['p', 'div']):
                if paragraph.find_parent(['ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    continue
                
                text = paragraph.get_text().strip()
                if text:
                    # Preserve paragraph numbering if present
                    numbered_match = re.match(r'^(\d+\.?\d*|\(\d+\)|\w+\.)\s+(.+)$', text)
                    if numbered_match:
                        number, content = numbered_match.groups()
                        result_text.append(f"{number} {content}")
                    else:
                        result_text.append(text)
            
            # Handle tables
            for table in soup.find_all('table'):
                result_text.append("\n--- TABLE ---\n")
                for row in table.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    result_text.append(" | ".join(cells))
                result_text.append("--- END TABLE ---\n")
            
            # Combine and clean
            combined_text = "\n\n".join(result_text)
            cleaned_lines = []
            for line in combined_text.split('\n'):
                line = re.sub(r'\s+', ' ', line).strip()
                if line:
                    cleaned_lines.append(line)
            
            final_text = '\n'.join(cleaned_lines)
            logger.info(f"Extracted {len(final_text)} characters from URL: {url}")
            
            return final_text
            
        except ImportError:
            raise ImportError("BeautifulSoup4 is required for URL processing. Install with 'pip install beautifulsoup4'")
        except requests.RequestException as e:
            raise ValueError(f"Error fetching URL {url}: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error processing URL {url}: {str(e)}")
    
    def can_process(self, file_type: str) -> bool:
        """Check if can process URLs."""
        return file_type.lower() == 'url'

class FileProcessingService:
    """
    Main file processing service that handles multiple file types.
    """
    
    def __init__(self):
        """Initialize the file processing service."""
        self.processors = {
            'pdf': PDFProcessor(),
            'docx': DOCXProcessor(), 
            'html': HTMLProcessor(),
            'htm': HTMLProcessor(),
            'txt': TextProcessor(),
            'text': TextProcessor(),
            'md': TextProcessor(),
            'markdown': TextProcessor(),
            'url': URLProcessor()
        }
    
    def process_file(self, file_path: str, file_type: str) -> str:
        """
        Process a file and extract its text content.
        
        Args:
            file_path: Path to the file (or URL if file_type is 'url')
            file_type: Type of the file ('pdf', 'docx', 'html', 'txt', 'url', etc.)
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported or processing fails
        """
        file_type = file_type.lower().strip()
        
        if file_type not in self.processors:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        processor = self.processors[file_type]
        
        try:
            return processor.process(file_path)
        except Exception as e:
            logger.error(f"Error processing {file_type} file {file_path}: {e}")
            raise
    
    def split_text(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        Split text into smaller chunks for processing.
        
        Args:
            text: Text to split
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        # Simple paragraph-based chunking
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size, store current chunk
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Include overlap from the end of previous chunk
                if len(current_chunk) > chunk_overlap:
                    current_chunk = current_chunk[-chunk_overlap:] + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def get_supported_types(self) -> List[str]:
        """
        Get list of supported file types.
        
        Returns:
            List of supported file type extensions
        """
        return list(self.processors.keys())
